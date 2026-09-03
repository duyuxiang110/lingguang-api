import os


def pdf_to_word_text(pdf_path: str, output_path: str) -> int:
    """文本模式：pdf2docx 高保真还原版式，失败时回退到纯文本提取。返回页数。"""
    try:
        return _pdf_to_word_layout(pdf_path, output_path)
    except Exception as e:
        print(f"[pdf2word] pdf2docx 转换失败，回退纯文本模式: {e}")
        return _pdf_to_word_plain(pdf_path, output_path)


def _pdf_to_word_layout(pdf_path: str, output_path: str) -> int:
    """高保真模式：pdf2docx 还原字体/版式/表格/图片。返回页数。"""
    import fitz
    from pdf2docx import Converter

    page_count = fitz.open(pdf_path).page_count
    # multi_processing=False：服务器内存有限，避免多进程占用翻倍
    cv = Converter(pdf_path)
    try:
        cv.convert(output_path, multi_processing=False)
    finally:
        cv.close()
    return page_count


def _pdf_to_word_plain(pdf_path: str, output_path: str) -> int:
    """纯文本兜底：pdfplumber 提取文本和表格 → python-docx 生成 Word。返回页数。"""
    import pdfplumber
    from docx import Document

    doc = Document()
    page_count = 0

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_count += 1
            tables = page.extract_tables()

            if tables:
                for table in tables:
                    if not table:
                        continue
                    rows = len(table)
                    cols = max(len(row) for row in table) if table else 0
                    if rows == 0 or cols == 0:
                        continue
                    doc_table = doc.add_table(rows=rows, cols=cols)
                    for i, row in enumerate(table):
                        for j, cell in enumerate(row):
                            if j < cols:
                                doc_table.rows[i].cells[j].text = cell or ""
                    doc.add_paragraph()
            else:
                text = page.extract_text() or ""
                if text.strip():
                    for line in text.split("\n"):
                        if line.strip():
                            doc.add_paragraph(line)
                    doc.add_paragraph()

            if page_count < len(pdf.pages):
                doc.add_page_break()

    doc.save(output_path)
    return page_count


def pdf_to_word_image(pdf_path: str, output_path: str, dpi: int = 200) -> int:
    """图片模式：PDF 每页渲染为图片嵌入 Word。返回页数。"""
    from pdf2image import convert_from_path
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    images = convert_from_path(pdf_path, dpi=dpi)

    for i, img in enumerate(images):
        temp_img = os.path.join(os.path.dirname(output_path), f"page_{i:04d}.png")
        img.save(temp_img, "PNG")
        doc.add_picture(temp_img, width=Inches(6.5))
        if i < len(images) - 1:
            doc.add_page_break()

    doc.save(output_path)

    for i in range(len(images)):
        temp_img = os.path.join(os.path.dirname(output_path), f"page_{i:04d}.png")
        if os.path.exists(temp_img):
            os.unlink(temp_img)

    return len(images)
